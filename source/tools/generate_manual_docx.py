from __future__ import annotations

from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]
DIST = ROOT / "dist"
DOCX_PATH = DIST / "scFlow Studio User Manual.docx"

MANUAL = """# scFlow Studio User Manual

## Overview
scFlow Studio is a graphical desktop application for single-cell RNA-seq analysis. It provides a guided workflow for data import, quality control, doublet removal, optional batch correction, clustering, annotation, subcluster analysis, differential expression, GSEA, single-gene visualization, gene-set scoring, and export.

## Standard Workflow
1. Create or open a project.
2. Add samples and verify input data.
3. Run single-sample QC.
4. Run or skip doublet removal.
5. Optionally run batch correction.
6. Merge samples and perform clustering.
7. Annotate main clusters.
8. Run subcluster analysis when needed.
9. Run differential expression and GSEA.
10. Generate single-gene and gene-set score visualizations.
11. Export figures, tables, h5ad files, project bundles, and Seurat objects.

## Input Formats
The application supports standard 10X matrix folders, compatible expression matrix files, archive imports, Seurat RDS import, marker templates, gene-set files, and local GMT pathway files.

## Output Formats
The application writes PNG, PDF, SVG, CSV/TSV/TXT, RDS, and h5ad outputs to the project cache and export directories.
"""


def main() -> None:
    DIST.mkdir(parents=True, exist_ok=True)
    try:
        from docx import Document
    except Exception:
        txt_path = DOCX_PATH.with_suffix(".md")
        txt_path.write_text(MANUAL, encoding="utf-8")
        print(f"python-docx is not installed. Wrote Markdown manual to: {txt_path}")
        return

    doc = Document()
    for line in MANUAL.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:], level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=2)
        elif line.strip() and line[0].isdigit() and ". " in line[:4]:
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
        elif line.startswith("- "):
            doc.add_paragraph(line[2:], style="List Bullet")
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(DOCX_PATH)
    print(f"Manual written to: {DOCX_PATH}")


if __name__ == "__main__":
    main()
